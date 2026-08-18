import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas

os.makedirs("docs", exist_ok=True)

# ---------------------------------------------------------------------------
# 1. MARKDOWN DOCUMENTATION GENERATION
# ---------------------------------------------------------------------------
md_content = """# Jumarald Pharmacy & Wellness — Comprehensive Technical & Operational Documentation

**Version:** 3.0.0 (Production Ready)  
**Superintendent Pharmacist:** Pharm. Philip Bruce-Tagoe (GPHC Reg. No. 2050984)  
**Compliance Standard:** Ghana Food & Drugs Authority (FDA), Pharmacy Council Ghana, Data Protection Act 2012 (Act 843)

---

## 1. Executive Summary

**Jumarald Pharmacy & Wellness** is a premier digital health and pharmaceutical enterprise application designed for Greater Accra and nationwide Ghana. The platform integrates e-commerce medicine ordering, prescription submission and verification workflows, telehealth doctor consultations, multi-branch inventory tracking, and an intelligent **Superintendent Clinical AI Assistant ("Dr. Jumarald AI")**.

The platform is engineered with modern web standards, prioritizing high visual aesthetics, strict data confidentiality, transaction integrity, automated safety fallbacks, and compliance with Ghanaian healthcare regulations.

---

## 2. System Architecture & Tech Stack

The solution is architected as a high-performance monorepo workspace:

| Layer | Technology / Tooling | Description |
| :--- | :--- | :--- |
| **Frontend Storefront** | Next.js 15 (App Router), React 19, TypeScript, Tailwind CSS, Framer Motion | Patient portal, medicine catalog, review submission, AI chat widget, telehealth booking |
| **Admin Console** | Next.js 15, React 19, TypeScript, Recharts, Lucide Icons | Pharmacy operational dashboard, order fulfillment, inventory management, review verification |
| **Backend API** | Express.js 4, Node.js (v18+), TypeScript, Socket.io | REST API, WebSockets real-time updates, AI orchestrator, Paystack payment webhooks |
| **Database & ORM** | Neon PostgreSQL (Serverless), Prisma ORM 5 | Relational data persistence, schema migrations, ACID transaction guarantees |
| **AI Triage Engine** | OpenAI GPT-4o-mini / Gemini 1.5 Flash + Fallback Engine | Multi-tool orchestrator, RAG knowledge search, interaction analysis, prescription decoding |
| **Storage & Media** | Cloudinary Cloud CDN / In-Memory Buffer | Encrypted prescription uploads, product image assets |
| **Security & Auth** | JWT Access Tokens, bcrypt, Zod, Helmet, Rate Limiter | Role-Based Access Control (RBAC), API rate limiting, CORS restriction |

---

## 3. Superintendent AI Assistant Architecture ("Dr. Jumarald AI")

The **Dr. Jumarald AI** subsystem functions as a virtual clinical pharmacy assistant operating under the core guiding principle: *"AI assists. Licensed Pharmacists decide."*

### 3.1 Dual-Provider Orchestration Engine
* **Provider Factory (`getAIProvider`)**: Dynamically selects between **OpenAI GPT-4o-mini** and **Google Gemini 1.5 Flash** based on configured API keys, providing automatic failover.
* **Intent Classification Pipeline (`AIOrchestrator`)**:
  1. `PRODUCT_SEARCH`: Multi-keyword search across master catalog with live pricing in GH₵ and stock validation.
  2. `DRUG_INTERACTION`: Pharmacological contraindication analysis across active ingredients.
  3. `PRESCRIPTION_ANALYSIS`: Decodes prescription shorthand (`TDS`, `BD`, `PRN`), dosage, frequency, and precautions.
  4. `BRANCH_INQUIRY`: Returns operating hours, contact numbers, and branch pickup availability.
  5. `SYMPTOM_TRIAGE`: Evaluates severity (`LOW`, `MODERATE`, `HIGH`, `URGENT`) and suggests doctor consultation.
  6. `GENERAL_CONSULTATION`: Plain-text empathetic clinical guidance sanitized of raw markdown headers (`##`/`###`).

### 3.2 Tool Execution Registry (`tool-registry.ts`)
* `searchProducts`: Multi-keyword catalog query returning live prices in GH₵ and availability.
* `getProductDetails`: Retrieves active ingredients, dosage guidance, and therapeutic indications.
* `checkDrugInteraction`: Multi-medication pharmacological contraindication analysis.
* `findNearbyBranches`: Returns branch physical addresses, phone numbers, and operating hours.
* `getUserOrders`: Authenticated patient order status tracking.
* `createPharmacistConsultation`: Escalates high-risk cases directly to the Superintendent Pharmacist queue (`AIEscalation`).

### 3.3 High-Reliability Clinical Fallback Engine (`fallbackClinicalEngine`)
* Zero-downtime protection ensuring that if third-party AI APIs experience latency or rate-limiting, the system automatically switches to an internal rule-based triage engine.
* Flags emergency symptoms (chest pain, acute shortness of breath) with immediate Ghana emergency warnings (**112 / 193**).

---

## 4. Comprehensive Security & CIA Triad Audit

### 4.1 Confidentiality Audit
* **Role-Based Access Control (RBAC)**: Enforced via `requireRole` middleware across sensitive routes (`SUPER_ADMIN`, `ADMIN`, `PHARMACIST`, `DOCTOR`, `PATIENT`).
* **Active User Database Verification**: `authenticateToken` middleware queries `prisma.user` on every request. Deactivated accounts return `403 Account disabled or session revoked`, invalidating active JWTs instantly.
* **Password Encryption**: Hashed using `bcrypt` (10-12 salt rounds). Raw passwords are never persisted or logged.
* **Prescription Data Privacy**: Prescriptions uploaded via memory buffers and encrypted Cloudinary streams with randomized public IDs.

### 4.2 Integrity Audit
* **ACID Transactions**: Multi-step inventory adjustments, order fulfillment, and loyalty point redemptions use `prisma.$transaction`.
* **Paystack HMAC SHA-512 Verification**: Paystack webhooks validate `x-paystack-signature` using HMAC SHA-512 before updating order statuses to `PAID`.
* **Strict Input Validation**: Endpoints process incoming data strictly validated against Zod schemas.
* **Audit Trail**: Administrative actions (stock transfers, bulk product edits, review approvals) recorded in the `AuditLog` table with user attribution and timestamps.

### 4.3 Availability Audit
* **Resource Rate Limiting**:
  * `authLimiter`: 20 login/register attempts per 15 minutes
  * `aiLimiter`: 20 requests per minute
  * `prescriptionLimiter`: 30 uploads per 15 minutes
  * `apiLimiter`: 120 general requests per minute
* **Strict Domain CORS**: Allowed origins restricted to Jumarald production subdomains (`jumarald*.vercel.app`, `jumaraldpharmacy.com`).
* **WebSocket Real-time Updates**: Real-time inventory sync and order status updates via Socket.io with auto-reconnection (`useSocketAutoInvalidate`).

---

## 5. Integrated External APIs & Services

| Provider / API | Purpose / Integration Point | Auth / Protocol |
| :--- | :--- | :--- |
| **Paystack Payment API** | Ghana Mobile Money (MTN MoMo, Telecel Cash, AT Money) & Card checkout in GH₵ | Bearer API Key & Webhook HMAC SHA-512 |
| **OpenAI API (GPT-4o-mini)** | Primary LLM engine for structured JSON tool execution & interaction analysis | Bearer API Key (`OPENAI_API_KEY`) |
| **Google Gemini API (1.5 Flash)** | Secondary LLM engine for clinical consultation & RAG knowledge search | API Key (`GEMINI_API_KEY`) |
| **Cloudinary Cloud CDN** | Encrypted prescription upload storage & product image optimization CDN | Cloud API Secret & Signed SDK Streams |
| **Nodemailer / SMTP API** | Transactional emails for order receipts, prescription verification, & password resets | TLS/SSL Authenticated SMTP |
| **WebPush API / VAPID Keys** | Real-time browser push notifications for rider tracking & restocking alerts | VAPID Key Pair Protocol |
| **Sentry Telemetry API** | Backend exception logging, performance monitoring, & error tracing | Sentry DSN & Node SDK |

---

## 6. Strategic Future Upgrades & Planned Purchases

1. **Twilio / Hubtel SMS Gateway API (Paid Enterprise Service)**: Direct SMS notifications for Ghanaian mobile numbers (instant OTP login codes, prescription status alerts, rider ETA updates).
2. **Upstash Redis / Dedicated Redis Cluster (Paid Subscription)**: Distributed in-memory session caching, high-performance rate limiting, and real-time Socket.io pub/sub state synchronization.
3. **AWS S3 / Cloudflare R2 Encrypted Health Vault (Paid Cloud Storage)**: Long-term HIPAA and Ghana Data Protection Act (Act 843) compliant encrypted prescription archiving.
4. **RxNorm / NLM Clinical Drug Interaction API (Enterprise License)**: Integration with international standardized pharmaceutical registries for automated contraindication scanning.
5. **IdentityPass / Ghana Card National ID Verification API (Paid Service)**: Automated identity validation for patient onboarding and doctor medical license validation.
6. **Google Maps Platform API (Paid Geocoding & Route Optimization)**: Exact delivery distance calculation, rider geolocation tracking, and dynamic delivery fee computation.

---

## 7. Regulatory & Legal Policies

1. **Acceptable Use Policy (`/acceptable-use`)**: Covers prescription fraud prevention, AI interaction rules, prohibited scraping, and account revocation rules.
2. **Terms of Service (`/terms`)**: Details Prescription-Only Medicine (POM) verification protocols, cold-chain transport regulations, non-returnable drug rules, and emergency medical disclaimers.
3. **Privacy Policy (`/privacy`)**: Details patient rights and data protection standards under the **Ghana Data Protection Act 2012 (Act 843)**.

---

## 8. Deployment & Maintenance

### Environment Variables (.env)
```env
PORT=5000
NODE_ENV=production
DATABASE_URL=postgresql://user:password@ep-neon-db.neon.tech/jumarald
JWT_SECRET=your_secure_jwt_secret_key
ALLOWED_ORIGINS=https://jumaraldpharmacy.com,https://jumaraldpharmacy.vercel.app
OPENAI_API_KEY=sk-proj-...
GEMINI_API_KEY=AIzaSy...
PAYSTACK_SECRET_KEY=sk_live_...
CLOUDINARY_CLOUD_NAME=jumarald
CLOUDINARY_API_KEY=...
CLOUDINARY_API_SECRET=...
```

### Build & Migration Commands
```bash
# Install dependencies
pnpm install

# Push Prisma Database Schema & Seed Data
cd backend
pnpm prisma:db-push
pnpm prisma:seed

# Production Build
pnpm build
```

---
© 2026 Jumarald Pharmacy & Wellness Ltd. All rights reserved.
"""

with open("docs/DOCUMENTATION.md", "w", encoding="utf-8") as f:
    f.write(md_content)

print("Generated docs/DOCUMENTATION.md successfully.")


# ---------------------------------------------------------------------------
# 2. MICROSOFT WORD DOCUMENT (.DOCX) GENERATION
# ---------------------------------------------------------------------------
def create_docx():
    doc = Document()
    
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    EMERALD = RGBColor(16, 185, 129)
    DARK_SLATE = RGBColor(15, 23, 42)
    GRAY = RGBColor(71, 85, 105)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = DARK_SLATE

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Jumarald Pharmacy & Wellness")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = EMERALD

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Full Technical Architecture, AI Engine & Security CIA Audit Documentation\nVersion 3.0.0")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.font.size = Pt(16)
        r.font.bold = True
        r.font.color.rgb = EMERALD
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.font.size = Pt(13)
        r.font.bold = True
        r.font.color.rgb = DARK_SLATE
        return h

    add_heading_1("1. Executive Summary")
    doc.add_paragraph(
        "Jumarald Pharmacy & Wellness is a full-stack digital health and pharmaceutical enterprise application operating in Ghana. "
        "The platform unifies e-commerce medicine retail, prescription validation, telehealth physician appointments, real-time multi-branch inventory tracking, "
        "and an AI-powered Superintendent Clinical Assistant ('Dr. Jumarald AI')."
    )

    add_heading_1("2. Superintendent AI Assistant Architecture ('Dr. Jumarald AI')")
    doc.add_paragraph(
        "Dr. Jumarald AI operates under the clinical rule: 'AI assists. Licensed Pharmacists decide.' "
        "It features a Dual-Provider Factory (OpenAI GPT-4o-mini / Gemini 1.5 Flash), an autonomous Intent Classifier, Tool Registry, and Rule-Based Fallback Engine."
    )
    
    add_heading_2("2.1 Intent Classifier Pipeline")
    doc.add_paragraph("• PRODUCT_SEARCH: Multi-keyword search across master catalog with live GH₵ prices and stock validation.")
    doc.add_paragraph("• DRUG_INTERACTION: Pharmacological contraindication analysis across active ingredients.")
    doc.add_paragraph("• PRESCRIPTION_ANALYSIS: Decodes prescription shorthand (TDS, BD, PRN), dosage, and precautions.")
    doc.add_paragraph("• BRANCH_INQUIRY: Returns branch physical addresses, phone numbers, and operating hours.")
    doc.add_paragraph("• SYMPTOM_TRIAGE: Evaluates severity (LOW, MODERATE, HIGH, URGENT) and suggests doctor consultation.")

    add_heading_2("2.2 Registered AI Tools")
    tools = [
        ("searchProducts", "Multi-keyword catalog query returning live prices in GH₵ and availability."),
        ("getProductDetails", "Retrieves active ingredients, dosage guidance, and indications."),
        ("checkDrugInteraction", "Multi-medication contraindication analysis."),
        ("findNearbyBranches", "Returns branch addresses, phone numbers, and operating hours."),
        ("getUserOrders", "Retrieves patient order tracking status for authenticated accounts."),
        ("createPharmacistConsultation", "Escalates high-risk cases directly to the Superintendent Pharmacist queue."),
    ]
    for name, desc in tools:
        p = doc.add_paragraph(style='List Bullet')
        r_name = p.add_run(f"{name}: ")
        r_name.bold = True
        p.add_run(desc)

    add_heading_1("3. Comprehensive Security & CIA Triad Audit")
    add_heading_2("3.1 Confidentiality Audit")
    doc.add_paragraph("Stateless JWT access tokens with RBAC middleware (requireRole). Database active-user validation on every token request revokes access immediately upon account deactivation. Passwords hashed using bcrypt (10-12 salt rounds). Encrypted Cloudinary streams for prescription documents.")
    
    add_heading_2("3.2 Integrity Audit")
    doc.add_paragraph("Prisma ACID transactions (prisma.$transaction) for inventory adjustments. Paystack webhooks validate x-paystack-signature using HMAC SHA-512. Input payloads sanitized using Zod schemas. Administrative actions recorded in AuditLog table.")

    add_heading_2("3.3 Availability Audit")
    doc.add_paragraph("Rate limiters protect system resources:")
    doc.add_paragraph("• authLimiter: 20 login attempts / 15 mins\n• aiLimiter: 20 requests / min\n• prescriptionLimiter: 30 uploads / 15 mins\n• apiLimiter: 120 general requests / min")
    doc.add_paragraph("Allowed origins restricted to Jumarald production subdomains. WebSocket real-time sync via Socket.io with auto-reconnection.")

    add_heading_1("4. Integrated External APIs")
    apis = [
        ("Paystack Payment API", "Ghana Mobile Money (MTN MoMo, Telecel, AT Money) & Card checkout in GH₵."),
        ("OpenAI GPT-4o-mini API", "Primary LLM for structured JSON tool execution and drug interaction analysis."),
        ("Google Gemini 1.5 Flash API", "Secondary LLM engine for clinical consultation and RAG knowledge search."),
        ("Cloudinary Cloud CDN API", "Encrypted prescription document storage and product image CDN."),
        ("Nodemailer / SMTP Mail Server", "Transactional emails for order receipts and password reset OTPs."),
        ("WebPush / VAPID Protocol", "Real-time browser notifications for rider tracking and inventory alerts."),
    ]
    for name, desc in apis:
        p = doc.add_paragraph(style='List Bullet')
        r_name = p.add_run(f"{name}: ")
        r_name.bold = True
        p.add_run(desc)

    add_heading_1("5. Future Software Procurement Roadmap")
    upgrades = [
        ("Twilio / Hubtel SMS Gateway (Paid API)", "Direct SMS alerts (OTP verification, prescription refill alerts, rider ETA)."),
        ("Upstash Redis Cluster (Paid Subscription)", "High-speed distributed in-memory caching and Socket.io pub/sub synchronization."),
        ("AWS S3 / Cloudflare R2 Vault (Paid Cloud)", "HIPAA and Ghana DPC compliant long-term encrypted prescription archiving."),
        ("RxNorm / NLM Drug Interaction API (License)", "Global pharmaceutical API integration for automated contraindication scanning."),
        ("IdentityPass Ghana Card Verification (Paid API)", "Automated identity validation for patients and doctors."),
        ("Google Maps Platform API (Paid License)", "Precise distance calculation, rider geolocation tracking, & dynamic delivery fees."),
    ]
    for name, desc in upgrades:
        p = doc.add_paragraph(style='List Bullet')
        r_name = p.add_run(f"{name}: ")
        r_name.bold = True
        p.add_run(desc)

    doc.save("docs/JUMARALD_PHARMACY_DOCUMENTATION.docx")
    print("Generated docs/JUMARALD_PHARMACY_DOCUMENTATION.docx successfully.")

create_docx()


# ---------------------------------------------------------------------------
# 3. PDF DOCUMENT GENERATION (ReportLab)
# ---------------------------------------------------------------------------
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#64748B"))
        
        if self._pageNumber > 1:
            self.drawString(54, 750, "Jumarald Pharmacy & Wellness — Technical Documentation")
            self.setStrokeColor(colors.HexColor("#E2E8F0"))
            self.setLineWidth(0.5)
            self.line(54, 742, 558, 742)

        self.setStrokeColor(colors.HexColor("#E2E8F0"))
        self.setLineWidth(0.5)
        self.line(54, 50, 558, 50)
        
        page_text = f"Page {self._pageNumber} of {page_count}"
        self.drawRightString(558, 36, page_text)
        self.drawString(54, 36, "© 2026 Jumarald Pharmacy & Wellness Ltd. · Confidential")
        self.restoreState()


def create_pdf():
    pdf_filename = "docs/JUMARALD_PHARMACY_DOCUMENTATION.pdf"
    doc = SimpleDocTemplate(
        pdf_filename,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=64
    )

    styles = getSampleStyleSheet()
    
    COLOR_PRIMARY = colors.HexColor("#10B981")
    COLOR_DARK = colors.HexColor("#0F172A")
    COLOR_MUTED = colors.HexColor("#475569")

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=COLOR_PRIMARY,
        alignment=1,
        spaceAfter=6
    )

    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=11,
        leading=15,
        textColor=COLOR_MUTED,
        alignment=1,
        spaceAfter=18
    )

    h1_style = ParagraphStyle(
        'Heading1_Custom',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=14,
        leading=17,
        textColor=COLOR_PRIMARY,
        spaceBefore=12,
        spaceAfter=5,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'Body_Custom',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=9,
        leading=13,
        textColor=COLOR_DARK,
        spaceAfter=5
    )

    bullet_style = ParagraphStyle(
        'Bullet_Custom',
        parent=body_style,
        leftIndent=12,
        firstLineIndent=-8,
        spaceAfter=3
    )

    story = []

    story.append(Spacer(1, 10))
    story.append(Paragraph("Jumarald Pharmacy & Wellness", title_style))
    story.append(Paragraph("Full AI Architecture & Security CIA Audit Documentation — v3.0.0", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=COLOR_PRIMARY, spaceBefore=0, spaceAfter=12))

    story.append(Paragraph("1. Executive Summary", h1_style))
    story.append(Paragraph(
        "<b>Jumarald Pharmacy & Wellness</b> is a digital healthcare platform operating in Ghana. "
        "The system unifies e-commerce pharmaceutical sales, certified prescription validation, telehealth physician appointments, "
        "real-time multi-branch inventory tracking, and an intelligent <b>Superintendent Clinical AI Assistant ('Dr. Jumarald AI')</b>.<br/>"
        "<b>Superintendent Pharmacist:</b> Pharm. Philip Bruce-Tagoe (GPHC Reg. No. 2050984).",
        body_style
    ))

    story.append(Paragraph("2. Superintendent AI Assistant Architecture ('Dr. Jumarald AI')", h1_style))
    story.append(Paragraph(
        "Dr. Jumarald AI features a Dual-Provider Factory (OpenAI GPT-4o-mini / Gemini 1.5 Flash), an autonomous Intent Classifier, Tool Registry, and Rule-Based Fallback Engine:",
        body_style
    ))
    story.append(Paragraph("• <b>Intent Classifier Pipeline:</b> PRODUCT_SEARCH, DRUG_INTERACTION, PRESCRIPTION_ANALYSIS, BRANCH_INQUIRY, SYMPTOM_TRIAGE, GENERAL_CONSULTATION.", bullet_style))
    story.append(Paragraph("• <b>Registered AI Tools:</b> searchProducts, getProductDetails, checkDrugInteraction, findNearbyBranches, getUserOrders, createPharmacistConsultation.", bullet_style))
    story.append(Paragraph("• <b>Clinical Fallback Engine:</b> Zero-downtime safety net triggering emergency warnings (Ghana 112 / 193).", bullet_style))

    story.append(Paragraph("3. Comprehensive Security & CIA Triad Audit", h1_style))
    story.append(Paragraph("<b>3.1 Confidentiality Audit:</b> Stateless JWT sessions, bcrypt password hashing, active user DB checks in <code>authenticateToken</code> middleware, encrypted Cloudinary streams.", body_style))
    story.append(Paragraph("<b>3.2 Integrity Audit:</b> Prisma ACID transactions for inventory updates, Zod input validation, Paystack HMAC SHA-512 webhook verification, AuditLog tracking.", body_style))
    story.append(Paragraph("<b>3.3 Availability Audit:</b> Dedicated rate limiters (<code>authLimiter</code>: 20/15m, <code>aiLimiter</code>: 20/m, <code>prescriptionLimiter</code>: 30/15m), domain-restricted CORS, Socket.io auto-invalidation.", body_style))

    story.append(Paragraph("4. Integrated External APIs & Services", h1_style))
    apis = [
        "<b>Paystack Payment API:</b> Ghana Mobile Money (MoMo) & Card checkout in GH₵ with HMAC SHA-512 webhooks.",
        "<b>OpenAI GPT-4o-mini & Gemini 1.5 Flash APIs:</b> Core LLM engines powering Dr. Jumarald AI.",
        "<b>Cloudinary Cloud CDN API:</b> Encrypted prescription uploads & optimized product images.",
        "<b>Nodemailer / SMTP API:</b> Transactional emails for receipts, prescription verification, & password resets.",
        "<b>WebPush / VAPID API:</b> Real-time browser notifications for rider tracking.",
        "<b>Sentry Telemetry API:</b> Application performance monitoring and backend error tracking."
    ]
    for api in apis:
        story.append(Paragraph(f"• {api}", bullet_style))

    story.append(Paragraph("5. Strategic Future Software Procurement Roadmap", h1_style))
    upgrades = [
        "<b>Twilio / Hubtel SMS Gateway API (Paid Service):</b> Direct SMS notifications for Ghanaian mobile numbers (OTP logins, instant prescription refill alerts, rider ETA).",
        "<b>Upstash Redis Cluster (Paid Subscription):</b> High-speed distributed in-memory session caching & Socket.io pub/sub synchronization.",
        "<b>AWS S3 / Cloudflare R2 Vault (Paid Cloud):</b> HIPAA & Ghana DPC compliant long-term encrypted prescription archiving.",
        "<b>RxNorm / NLM Drug Interaction API (License):</b> Global pharmaceutical API integration for automated contraindication scanning.",
        "<b>IdentityPass Ghana Card Verification API (Paid Service):</b> Automated identity validation for patients and doctors.",
        "<b>Google Maps Platform API (Paid License):</b> Precise distance calculation, rider geolocation tracking, & dynamic delivery fees."
    ]
    for u in upgrades:
        story.append(Paragraph(f"• {u}", bullet_style))

    doc.build(story, canvasmaker=NumberedCanvas)
    print("Generated docs/JUMARALD_PHARMACY_DOCUMENTATION.pdf successfully.")

create_pdf()
